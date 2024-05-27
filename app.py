import os
from flask import Flask, request, render_template, send_file, after_this_request
import pandas as pd
from docx import Document
from docx.shared import Pt
import tempfile
import zipfile

app = Flask(__name__)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        regional_selecionada = request.form['regional']
        file = request.files['file']
        filename = file.filename
        tempdir = tempfile.mkdtemp()
        filepath = os.path.join(tempdir, filename)
        file.save(filepath)

        termos_dir = os.path.join(tempdir, 'termos')
        os.makedirs(termos_dir, exist_ok=True)

        gerar_termos(filepath, regional_selecionada, termos_dir)

        zip_filename = os.path.join(tempdir, 'termos.zip')
        with zipfile.ZipFile(zip_filename, 'w') as zipf:
            for root, _, files in os.walk(termos_dir):
                for file in files:
                    zipf.write(os.path.join(root, file),
                               os.path.relpath(os.path.join(root, file),
                                               termos_dir))

        @after_this_request
        def cleanup(response):
            try:
                os.remove(zip_filename)
                os.removedirs(tempdir)
            except Exception as e:
                print(f"Error cleaning up: {e}")
            return response

        return send_file(zip_filename, as_attachment=True)

    return render_template('index.html')


def gerar_termos(arquivo_xlsx, regional_selecionada, pasta_destino):
    df = pd.read_excel(arquivo_xlsx)
    df_regional = df[df['REGIONAL'] == regional_selecionada]

    template_path = os.path.join(os.path.dirname(__file__), 'TERMO.docx')

    for index, row in df_regional.iterrows():
        doc = Document(template_path)

        for paragraph in doc.paragraphs:
            if 'nomeespecifico' in paragraph.text:
                paragraph.text = paragraph.text.replace('nomeespecifico', f"{row['SERVIDOR']}")
            if 'cpfespecifico' in paragraph.text:
                paragraph.text = paragraph.text.replace('cpfespecifico', f"{row['CPF']}")
            if 'matriculaespecifica' in paragraph.text:
                paragraph.text = f"Número funcional (matrícula):  {row['MATRICULA']}"
            if 'cargospecifico' in paragraph.text:
                paragraph.text = f"Cargo: {row['CARGO']}"
            if 'unidadespecifico' in paragraph.text:
                paragraph.text = f"Unidade de lotação: {row['UNIDADE']}"
            if 'regionalespecifica' in paragraph.text:
                paragraph.text = f"Regional: Superintendência Regional de Educação de {row['REGIONAL']}"

        def update_font_size(paragraph, font_size):
            for run in paragraph.runs:
                font = run.font
                font.size = Pt(font_size)

        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for paragraph in cell.paragraphs:
                        if 'codigoplaqueta' in paragraph.text:
                            if 'NUM. PLAQUETA' in row and pd.notna(row['NUM. PLAQUETA']):
                                num_plaqueta = str(row['NUM. PLAQUETA'])
                                paragraph.text = paragraph.text.replace('codigoplaqueta', f"{num_plaqueta}")
                            else:
                                paragraph.text = paragraph.text.replace('codigoplaqueta', '')
                            update_font_size(paragraph, 10)
                        if 'codigoserie' in paragraph.text:
                            if 'NUM. SERIE' in row and pd.notna(row['NUM. SERIE']):
                                num_serie = str(row['NUM. SERIE'])
                                paragraph.text = paragraph.text.replace('codigoserie', f"{num_serie}")
                            else:
                                paragraph.text = paragraph.text.replace('codigoserie', '')
                            update_font_size(paragraph, 10)

        doc.save(os.path.join(pasta_destino, f"{row['UNIDADE']}_TERMO_{row['SERVIDOR']}.docx"))


if __name__ == '__main__':
    app.run(debug=True)
